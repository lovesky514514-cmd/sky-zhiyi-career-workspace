import os, json, requests
from io import BytesIO
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS

def load_env_file():
    env_path = os.path.join(os.path.dirname(__file__), ".env")
    if not os.path.exists(env_path):
        return
    with open(env_path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            k, v = line.split("=", 1)
            os.environ.setdefault(k.strip(), v.strip())

load_env_file()

app = Flask(__name__)
CORS(app)

DEEPSEEK_URL = os.getenv("DEEPSEEK_BASE_URL", "https://api.deepseek.com/chat/completions")
DEEPSEEK_MODEL = os.getenv("DEEPSEEK_MODEL", "deepseek-chat")

def ask_deepseek(prompt, json_mode=True):
    api_key = os.getenv("DEEPSEEK_API_KEY", "").strip()
    if not api_key or api_key == "把你的DeepSeek_API_Key粘贴到这里":
        return None

    payload = {
        "model": DEEPSEEK_MODEL,
        "messages": [
            {"role": "system", "content": "你是一个面向大学生求职场景的AI助手。你必须只输出合法JSON，不要输出解释、Markdown或代码块。"},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.2,
        "stream": False
    }
    if json_mode:
        payload["response_format"] = {"type": "json_object"}

    resp = requests.post(
        DEEPSEEK_URL,
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
        timeout=60
    )
    resp.raise_for_status()
    text = resp.json()["choices"][0]["message"]["content"]

    try:
        return json.loads(text)
    except Exception:
        start, end = text.find("{"), text.rfind("}")
        if start >= 0 and end > start:
            return json.loads(text[start:end+1])
        raise RuntimeError("DeepSeek返回内容不是JSON：" + text[:300])

@app.get("/api/health")
def health():
    return jsonify({
        "ok": True,
        "model": DEEPSEEK_MODEL,
        "base_url": DEEPSEEK_URL,
        "has_key": bool(os.getenv("DEEPSEEK_API_KEY", "").strip())
    })

@app.post("/api/parse")
def parse_resume():
    text = request.json.get("text", "")
    prompt = f"""
你是“材料解析智能体”。请从大学生简历/经历材料中提取结构化求职信息。
要求：
1. 不要编造事实；
2. 目标岗位可以根据文本合理推断；
3. missing_fields 写出还需要用户补充的内容；
4. 只输出JSON。

输出JSON字段：
{{
  "education": "string",
  "target_jobs": ["string"],
  "skills": ["string"],
  "experiences": [
    {{"type":"项目/实习/竞赛/课程/论文/社团/其他", "title":"string", "summary":"string"}}
  ],
  "missing_fields": ["string"],
  "agent_source": "deepseek-chat"
}}

材料：
{text[:10000]}
"""
    data = ask_deepseek(prompt) or {
        "education": "待补充",
        "target_jobs": ["待探索岗位方向"],
        "skills": [],
        "experiences": [{"type":"原始经历","title":"用户导入材料","summary":text[:220]}],
        "missing_fields": ["量化成果", "目标岗位优先级", "每段经历中的本人职责"],
        "agent_source": "server_fallback"
    }
    return jsonify(data)

@app.post("/api/refine")
def refine_memory():
    parsed = request.json.get("parsed", {})
    prompt = f"""
你是“经历精炼智能体”。请将解析结果精炼为可长期复用的求职记忆。
要求：
1. 去重、归并、提炼能力标签；
2. 不要编造奖项、公司、数据；
3. 输出内容可直接用于岗位适配、简历优化和面试训练；
4. 只输出JSON。

输出JSON字段：
{{
  "memory_type": "career_asset",
  "title": "string",
  "refined_summary": "string",
  "ability_tags": ["string"],
  "job_relevance": ["string"],
  "importance_score": 0-100,
  "resume_value": "high/medium/low",
  "suggested_layer": "long_term/transition_memory/working_memory",
  "agent_source": "deepseek-chat"
}}

解析结果：
{json.dumps(parsed, ensure_ascii=False)[:10000]}
"""
    data = ask_deepseek(prompt) or {
        "memory_type": "career_asset",
        "title": "求职经历资产",
        "refined_summary": "已整理为可复用的求职记忆。",
        "ability_tags": [],
        "job_relevance": ["待探索岗位方向"],
        "importance_score": 80,
        "resume_value": "high",
        "suggested_layer": "long_term",
        "agent_source": "server_fallback"
    }
    return jsonify(data)

@app.post("/api/interview")
def interview():
    target = request.json.get("target", "目标岗位")
    memories = request.json.get("memories", [])
    prompt = f"""
你是“面试训练智能体”。请根据用户求职记忆和目标岗位生成个性化训练方案。
要求：
1. 不编造不存在的经历；
2. 自我介绍控制在1分钟左右；
3. 面试题要包含通用题、岗位题、项目追问题；
4. 只输出JSON。

输出JSON字段：
{{
  "self_intro": "string",
  "questions": [
    {{"question":"string", "why":"string", "answer_hint":"string"}}
  ]
}}

目标岗位：{target}
用户记忆：
{json.dumps(memories, ensure_ascii=False)[:10000]}
"""
    data = ask_deepseek(prompt) or {
        "self_intro": "请先补充更多求职记忆，系统会生成更个性化的自我介绍。",
        "questions": [
            {"question":"请简单介绍一下自己","why":"考查表达和岗位匹配","answer_hint":"按专业背景、核心经历、求职意向回答。"}
        ]
    }
    return jsonify(data)


@app.post("/api/evaluate")
def evaluate_answer():
    question = request.json.get("question", "")
    answer = request.json.get("answer", "")
    target = request.json.get("target", "目标岗位")
    prompt = f"""
你是面试反馈教练。请评价用户对面试问题的回答。
只输出JSON：
{{
  "overall": "总体评价",
  "strengths": ["优点"],
  "suggestions": ["改进建议"],
  "polished_answer": "优化后的示范回答"
}}
目标岗位：{target}
面试问题：{question}
用户回答：{answer}
"""
    data = ask_deepseek(prompt) or {
        "overall": "回答已收到，可继续补充更具体的经历、数据和结果。",
        "strengths": ["有基本回答方向"],
        "suggestions": ["补充STAR结构", "突出本人职责", "加入量化结果"],
        "polished_answer": answer
    }
    return jsonify(data)



JOBS_CACHE = None

def load_jobs():
    global JOBS_CACHE
    if JOBS_CACHE is not None:
        return JOBS_CACHE
    data_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "data", "jobs.json"))
    try:
        with open(data_path, "r", encoding="utf-8") as f:
            JOBS_CACHE = json.load(f)
    except Exception:
        JOBS_CACHE = []
    return JOBS_CACHE

def unique_items(items):
    out = []
    seen = set()
    for x in items:
        if not x:
            continue
        if x not in seen:
            out.append(x)
            seen.add(x)
    return out

def tokenize_query(q):
    q = (q or "").lower().replace("，", " ").replace(",", " ").replace("/", " ").replace("、", " ")
    tokens = [x.strip() for x in q.split() if x.strip()]
    raw = q
    if "ai产品" in raw:
        tokens += ["ai", "产品", "产品助理"]
    if "数据" in raw:
        tokens += ["数据", "分析", "数据分析"]
    if "材料" in raw:
        tokens += ["材料", "研发"]
    if "科研" in raw:
        tokens += ["科研", "助理"]
    if "运营" in raw:
        tokens += ["运营", "用户"]
    return unique_items(tokens)

def score_job_local(job, keyword="", city=""):
    text = " ".join(str(v) for v in job.values()).lower()
    job_name = str(job.get("job", "")).lower()
    industry = str(job.get("industry", "")).lower()
    company = str(job.get("company", "")).lower()
    job_city = str(job.get("city", ""))
    tokens = tokenize_query(keyword)
    score = 22
    reasons = []
    if not tokens and not city:
        score += 18

    for t in tokens:
        if t in job_name:
            score += 26
            reasons.append(f"岗位相关：{t}")
        elif t in industry:
            score += 14
            reasons.append(f"行业相关：{t}")
        elif t in company:
            score += 10
            reasons.append(f"单位相关：{t}")
        elif t in text:
            score += 7
            reasons.append(f"信息相关：{t}")

    # 通用方向映射，减少“产品助理”这类重复长句
    if "产品" in job_name or "产品" in industry:
        score += 8
    if "数据" in job_name or "分析" in job_name:
        score += 8
    if "材料" in job_name or "研发" in job_name:
        score += 8
    if "实习" in job_name or "实习" in str(job.get("type", "")):
        score += 4

    if city:
        if city in job_city:
            score += 18
            reasons.append(f"城市匹配：{city}")
        else:
            score -= 30

    if "否" in str(job.get("exam", "")):
        score += 4
        reasons.append("免笔试")

    if not tokens:
        reasons.append("城市通用推荐" if city else "方向待定，先看通用机会")
    if not reasons:
        reasons.append("可作为备选")

    return min(max(score, 0), 98), " / ".join(unique_items(reasons)[:3])

@app.post("/api/jobs")
def search_jobs():
    payload = request.json or {}
    keyword = payload.get("keyword", "")
    city = payload.get("city", "")
    page = int(payload.get("page", 1) or 1)
    page_size = int(payload.get("page_size", 30) or 30)
    page = max(page, 1)
    page_size = min(max(page_size, 1), 100)

    jobs = load_jobs()
    scored = []
    for job in jobs:
        if city and city not in str(job.get("city", "")):
            continue
        score, reasons = score_job_local(job, keyword, city)
        threshold = 50 if (keyword or city) else 35
        if score < threshold:
            continue
        item = dict(job)
        item["score"] = score
        item["reasons"] = reasons
        scored.append(item)

    scored.sort(key=lambda x: x.get("score", 0), reverse=True)
    total = len(scored)
    start = (page - 1) * page_size
    end = start + page_size

    return jsonify({
        "total": total,
        "page": page,
        "page_size": page_size,
        "rows": scored[start:end],
        "source": "backend_local_jobs_json"
    })



def extract_text_from_file(file_storage):
    filename = (file_storage.filename or "").lower()
    data = file_storage.read()

    if filename.endswith(".txt") or filename.endswith(".md"):
        for enc in ["utf-8", "gbk", "gb18030"]:
            try:
                return data.decode(enc)
            except Exception:
                pass
        return data.decode("utf-8", errors="ignore")

    if filename.endswith(".pdf"):
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(BytesIO(data))
            return "\n".join(page.extract_text() or "" for page in reader.pages).strip()
        except Exception as e:
            return f"PDF解析失败：{str(e)}。请复制正文粘贴。"

    if filename.endswith(".docx"):
        try:
            from docx import Document
            doc = Document(BytesIO(data))
            return "\n".join(p.text for p in doc.paragraphs).strip()
        except Exception as e:
            return f"DOCX解析失败：{str(e)}。请复制正文粘贴。"

    return "暂不支持该文件格式，请上传 txt、md、pdf 或 docx。"

@app.post("/api/upload_resume")
def upload_resume():
    if "file" not in request.files:
        return jsonify({"ok": False, "message": "未接收到文件。", "text": ""}), 400
    file = request.files["file"]
    text = extract_text_from_file(file)
    return jsonify({
        "ok": True,
        "filename": file.filename,
        "text": text
    })



@app.post("/api/resume-builder/generate")
def resume_builder_generate():
    payload = request.json or {}
    template = payload.get("template", "cn")
    target = payload.get("target_position", "目标岗位")
    city = payload.get("city", "")
    jd = payload.get("jd", "")
    materials = payload.get("materials", "")
    parsed = payload.get("parsed") or {}
    memories = payload.get("memories") or []
    is_en = template == "en"

    prompt = f"""
你是 SKY职忆的智能简历生成器。请根据用户材料、岗位JD和求职记忆，生成一份{'英文' if is_en else '中文'}简历的结构化JSON。
严格要求：
1. 只使用用户材料中出现过的信息，不要编造学校、公司、奖项、证书、年份或量化数据；
2. 如果缺少姓名/电话/邮箱，用“待补充”或英文 "To be added"；
3. 经历描述必须适合投递 {target}；
4. 每段经历用2-4条 bullets，突出本人职责、动作、结果；
5. 输出必须是合法JSON，不要Markdown。

输出JSON字段：
{{
  "language": "{template}",
  "name": "string",
  "target_position": "string",
  "phone": "string",
  "email": "string",
  "location": "string",
  "education": ["string"],
  "experience": [{{"title":"string", "role":"string", "time":"string", "bullets":["string"]}}],
  "projects": [{{"title":"string", "role":"string", "time":"string", "bullets":["string"]}}],
  "skills": ["string"],
  "certificates": ["string"],
  "self_summary": "string",
  "innovation_report": {{
    "skill_capital": 0-100,
    "job_fit": 0-100,
    "novelty_score": 0-100,
    "repeat_risk": 0-100,
    "suggestions": ["string"]
  }}
}}

目标岗位：{target}
目标城市：{city}
岗位JD：{jd[:5000]}
解析结果：{json.dumps(parsed, ensure_ascii=False)[:6000]}
求职记忆：{json.dumps(memories, ensure_ascii=False)[:8000]}
用户补充材料：{materials[:10000]}
"""
    data = ask_deepseek(prompt)
    if data:
        return jsonify(data)

    # fallback，保证无API key时仍可演示
    skills = []
    if isinstance(parsed, dict):
        skills.extend(parsed.get("skills") or [])
    for m in memories:
        skills.extend(m.get("ability_tags") or [])
    skills = unique_items([str(x) for x in skills])[:12]
    if not skills:
        skills = ["快速学习", "沟通表达", "项目执行"] if not is_en else ["Fast learning", "Communication", "Project execution"]
    exp_summary = ""
    if isinstance(parsed, dict) and parsed.get("experiences"):
        exp_summary = parsed.get("experiences", [{}])[0].get("summary", "")
    if not exp_summary and memories:
        exp_summary = memories[0].get("refined_summary", "")
    fallback = {
        "language": template,
        "name": "To be added" if is_en else "姓名待补充",
        "target_position": target,
        "phone": "To be added" if is_en else "待补充",
        "email": "To be added" if is_en else "待补充",
        "location": city or ("China" if is_en else "求职城市待补充"),
        "education": [parsed.get("education", "Bachelor Candidate") if isinstance(parsed, dict) else "Bachelor Candidate"],
        "experience": [{
            "title": target,
            "role": "Candidate" if is_en else "候选人",
            "time": "",
            "bullets": [exp_summary or ("Organized personal experience into structured resume materials." if is_en else "将个人经历整理为适合岗位投递的结构化简历材料。")]
        }],
        "projects": [],
        "skills": skills,
        "certificates": [],
        "self_summary": f"A motivated candidate for {target}, with strong learning ability and project execution." if is_en else f"面向{target}方向，具备快速学习、材料整理、沟通表达和项目执行能力。",
        "innovation_report": {
            "skill_capital": min(96, 45 + len(skills) * 5),
            "job_fit": 72 if jd else 64,
            "novelty_score": 78,
            "repeat_risk": 22,
            "suggestions": ["建议补充真实姓名、联系方式、时间线和量化成果。" if not is_en else "Add real contact information, timeline and quantified outcomes."]
        }
    }
    return jsonify(fallback)


def _as_list(v):
    if not v:
        return []
    if isinstance(v, list):
        return v
    return [v]


def _item_text(item):
    if isinstance(item, str):
        return item
    if not isinstance(item, dict):
        return str(item)
    title = item.get("title") or item.get("school") or item.get("company") or item.get("name") or "经历"
    role = item.get("role") or item.get("major") or item.get("position") or ""
    time = item.get("time") or ""
    prefix = "  ".join(x for x in [time, title, role] if x)
    bullets = item.get("bullets") or ([item.get("summary")] if item.get("summary") else [])
    if bullets:
        return prefix + "\n" + "\n".join(["• " + str(b) for b in bullets if b])
    return prefix


def build_resume_docx(draft, template="cn"):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_TABLE_ALIGNMENT

    is_en = template == "en"
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Pt(36)
    sec.bottom_margin = Pt(36)
    sec.left_margin = Pt(42)
    sec.right_margin = Pt(42)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(9.5)

    name = draft.get("name") or ("Your Name" if is_en else "姓名")
    title = draft.get("target_position") or draft.get("objective") or ""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(17, 24, 39)

    info = " | ".join([x for x in [title, draft.get("phone"), draft.get("email"), draft.get("location")] if x])
    p = doc.add_paragraph(info)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(107, 114, 128)

    def heading(txt):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(txt)
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(37, 99, 235)
        return p

    def bullet(txt):
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Pt(12)
        p.paragraph_format.space_after = Pt(2)
        p.add_run("• ").bold = True
        p.add_run(str(txt))

    heading("Objective" if is_en else "求职定位")
    doc.add_paragraph(draft.get("self_summary") or "")

    heading("Education" if is_en else "教育背景")
    for item in _as_list(draft.get("education")):
        doc.add_paragraph(_item_text(item))

    heading("Experience" if is_en else "项目 / 实习经历")
    for item in _as_list(draft.get("experience")) + _as_list(draft.get("projects")):
        if isinstance(item, dict):
            p = doc.add_paragraph()
            r = p.add_run("  ".join([x for x in [item.get("time"), item.get("title"), item.get("role")] if x]))
            r.bold = True
            for b in item.get("bullets") or []:
                bullet(b)
        else:
            doc.add_paragraph(str(item))

    heading("Skills & Certificates" if is_en else "技能证书")
    skills = _as_list(draft.get("skills")) + _as_list(draft.get("certificates"))
    doc.add_paragraph("；".join([str(x) for x in skills]))

    report = draft.get("innovation_report") or {}
    if report:
        heading("AI Resume Check" if is_en else "AI 简历生成检查")
        doc.add_paragraph(f"Skill Capital: {report.get('skill_capital', '--')} | Job Fit: {report.get('job_fit', '--')} | Novelty: {report.get('novelty_score', '--')} | Repeat Risk: {report.get('repeat_risk', '--')}%")
        for s in report.get("suggestions") or []:
            bullet(s)
    return doc


@app.post("/api/resume-builder/export")
def resume_builder_export():
    payload = request.json or {}
    template = payload.get("template", "cn")
    draft = payload.get("draft") or {}
    doc = build_resume_docx(draft, template)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    filename = "SKY_ZhiYi_Resume_EN.docx" if template == "en" else "SKY_ZhiYi_Resume_CN.docx"
    return send_file(bio, as_attachment=True, download_name=filename, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


if __name__ == "__main__":
    print("DeepSeek proxy server running.")
    print("Model:", DEEPSEEK_MODEL)
    print("URL:", DEEPSEEK_URL)
    app.run(host="127.0.0.1", port=5000, debug=True)
